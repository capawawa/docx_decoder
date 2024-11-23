import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Step 1: Convert the editable JSON-like format back to a .docx file with the original formatting intact

def editable_format_to_docx(json_path, output_docx_path):
    with open(json_path, 'r') as f:
        doc_data = json.load(f)

    doc = Document()

    # Restore header and footer if present
    if 'header' in doc_data[0]:
        for section in doc.sections:
            header = section.header.paragraphs[0]
            header.text = doc_data[0]['header']
            footer = section.footer.paragraphs[0]
            footer.text = doc_data[0]['footer']
        doc_data = doc_data[1:]

    for para in doc_data:
        p = doc.add_paragraph(para['text'])
        p.style = para['style']

        # Restore formatting details
        if para['alignment'] is not None:
            p.alignment = para['alignment']
        if para['space_before'] is not None:
            p.paragraph_format.space_before = Pt(para['space_before'])
        if para['space_after'] is not None:
            p.paragraph_format.space_after = Pt(para['space_after'])
        if para['left_indent'] is not None:
            p.paragraph_format.left_indent = Pt(para['left_indent'])
        if para['right_indent'] is not None:
            p.paragraph_format.right_indent = Pt(para['right_indent'])
        if para['first_line_indent'] is not None:
            p.paragraph_format.first_line_indent = Pt(para['first_line_indent'])
        if para['line_spacing'] is not None:
            p.paragraph_format.line_spacing = para['line_spacing']
        if para['line_spacing_rule'] is not None:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING(para['line_spacing_rule'])
        if para['keep_with_next'] is not None:
            p.paragraph_format.keep_with_next = para['keep_with_next']
        if para['keep_together'] is not None:
            p.paragraph_format.keep_together = para['keep_together']
        if para['page_break_before'] is not None:
            p.paragraph_format.page_break_before = para['page_break_before']
        if para['is_bullet']:
            p.style = 'List Bullet'
        if para['is_numbered']:
            p.style = 'List Number'

        # Restore run formatting (font properties)
        if para['font_name'] or para['font_size'] or para['font_bold'] or para['font_italic'] or para['font_underline'] or para['font_highlight_color'] or para['font_color']['rgb']:
            run = p.runs[0] if p.runs else p.add_run()
            if para['font_name']:
                run.font.name = para['font_name']
            if para['font_size']:
                run.font.size = Pt(para['font_size'])
            if para['font_bold'] is not None:
                run.font.bold = para['font_bold']
            if para['font_italic'] is not None:
                run.font.italic = para['font_italic']
            if para['font_underline'] is not None:
                run.font.underline = para['font_underline']
            if para['font_highlight_color'] is not None:
                highlight = OxmlElement('w:highlight')
                highlight.set(qn('w:val'), para['font_highlight_color'])
                run._r.get_or_add_pPr().append(highlight)
            if para['font_color']['rgb']:
                run.font.color.rgb = RGBColor.from_string(para['font_color']['rgb'])

    doc.save(output_docx_path)
    print(f"Document saved as '{output_docx_path}'")

if __name__ == "__main__":
    # Example usage
    json_path = 'resume_editable.json'
    output_docx_path = 'Restored_Resume.docx'

    # Step 1: Convert the edited JSON back to .docx
    editable_format_to_docx(json_path, output_docx_path)
