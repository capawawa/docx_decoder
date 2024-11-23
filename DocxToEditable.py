from docx import Document
import json

# Step 1: Read the .docx file and convert it to an editable JSON-like format
def docx_to_editable_format(docx_path):
    doc = Document(docx_path)
    doc_data = []

    for para in doc.paragraphs:
        doc_data.append({
            'text': para.text,
            'style': para.style.name
        })

    # Store it in a JSON file for easy editing later
    with open('resume_editable.json', 'w') as f:
        json.dump(doc_data, f, indent=4)

    print("Document converted to editable format and saved as 'resume_editable.json'")

# Step 2: Edit the JSON file manually or programmatically as needed
# This function allows you to make changes, but you can also directly edit the JSON file.
def edit_resume_json(data):
    # Example edit: Change the job title
    for item in data:
        if 'Brand Ambassador' in item['text']:
            item['text'] = 'Sales Account Manager'

    with open('resume_editable.json', 'w') as f:
        json.dump(data, f, indent=4)

    print("Changes saved to 'resume_editable.json'")

# Step 3: Convert JSON back to .docx with the original formatting intact
def editable_format_to_docx(json_path, output_docx_path):
    with open(json_path, 'r') as f:
        doc_data = json.load(f)

    doc = Document()
    for para in doc_data:
        p = doc.add_paragraph(para['text'])
        p.style = para['style']

    doc.save(output_docx_path)
    print(f"Document saved as '{output_docx_path}'")

if __name__ == "__main__":
    # Example usage:
    docx_path = 'Adam_Capuana_Resume.docx'
    json_path = 'resume_editable.json'
    output_docx_path = 'Updated_Adam_Capuana_Resume.docx'

    # Step 1: Convert the .docx to editable JSON-like format
    docx_to_editable_format(docx_path)

    # Step 2: Edit the JSON file manually or programmatically
    with open(json_path, 'r') as f:
        data = json.load(f)
    edit_resume_json(data)

    # Step 3: Convert the edited JSON back to .docx
    editable_format_to_docx(json_path, output_docx_path)
